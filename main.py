import os
import io
import re
from PIL import Image
from docx import Document
import win32com.client as win32
from docx.shared import Inches
from spellchecker import SpellChecker

def convert_doc_to_docx(doc_path):
    """Convert a DOC file to DOCX format."""
    word = win32.gencache.EnsureDispatch('Word.Application')
    doc = word.Documents.Open(doc_path)
    temp_docx_path = doc_path.replace(".doc", "_temp.docx")
    doc.SaveAs(temp_docx_path, FileFormat=16)  # 16 is for docx
    doc.Close()
    word.Quit()
    final_docx_path = doc_path.replace(".doc", ".docx")
    if os.path.exists(final_docx_path):
        os.remove(final_docx_path)
    os.rename(temp_docx_path, final_docx_path)
    return final_docx_path


def remove_protection(docx_path):
    """Remove protection from a DOCX file using win32com."""
    word = win32.gencache.EnsureDispatch('Word.Application')
    doc = word.Documents.Open(docx_path)
    if doc.ProtectionType != -1:  # Check if the document is protected
        doc.Unprotect()
    doc.Save()
    doc.Close()
    word.Quit()

def extract_images_from_doc(doc_path, temp_image_dir):
    """Extract all images from the DOCX file and save them to the temporary directory."""
    doc = Document(doc_path)
    if not os.path.exists(temp_image_dir):
        os.makedirs(temp_image_dir)

    image_count = 0
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            image_part = rel.target_part
            image_stream = io.BytesIO(image_part.blob)
            image = Image.open(image_stream)
            image_file_path = os.path.join(temp_image_dir, f'image_{image_count}.png')
            image.save(image_file_path)
            image_count += 1

    return image_count

def clean_text(text):
    """Remove extra spaces and line breaks between sentences in the text."""
    text = re.sub(r'\n+', '\n', text)
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'\s+([,.!?;:])', r'\1', text)
    text = re.sub(r'([,.!?;:])(\w)', r'\1 \2', text)
    return text

def correct_spelling(text):
    """Correct the spelling of the text."""
    spell = SpellChecker()
    corrected_text = []
    for word in text.split():
        corrected_word = spell.correction(word)
        if corrected_word is None:
            corrected_word = word
        corrected_text.append(corrected_word)
    return ' '.join(corrected_text)


def extract_vulnerabilities(doc):
    """Extract vulnerabilities from the tables in the DOCX file."""
    vulnerabilities = []
    for table in doc.tables:
        vulnerability = {}
        for row in table.rows:
            if "Vulnerability Name" in row.cells[0].text:
                vulnerability_name = row.cells[1].text.strip()
                vulnerability["name"] = correct_spelling(vulnerability_name)
            elif "Affected Host" in row.cells[0].text:
                affected_host = row.cells[1].text.strip()
                vulnerability["host"] = correct_spelling(affected_host)
            elif "Risk Rating" in row.cells[0].text:
                risk_rating = row.cells[1].text.strip()
                vulnerability["risk_rating"] = correct_spelling(risk_rating)

        if vulnerability:
            vulnerabilities.append(vulnerability)

    return vulnerabilities

def extract_content_between_headings(doc_path, start_heading, end_heading, output_path, temp_image_dir):
    """Extract content and images between specified headings and add table data."""
    doc = Document(doc_path)
    new_doc = Document()

    vulnerabilities = extract_vulnerabilities(doc)

    in_section = False
    vul_index = 0

    if not os.path.exists(temp_image_dir):
        os.makedirs(temp_image_dir)

    for para in doc.paragraphs:
        if para.style.name.startswith('Heading'):
            if para.text.strip() == start_heading:
                in_section = True
                continue
            elif para.text.strip() == end_heading:
                in_section = False
                continue

        if in_section:
            if para.text.strip() == "Description" and vul_index < len(vulnerabilities):
                vul = vulnerabilities[vul_index]
                new_doc.add_paragraph(f"Vulnerability Name: {vul.get('name', '')}")
                new_doc.add_paragraph(f"Affected Host: {vul.get('host', '')}")
                new_doc.add_paragraph(f"Risk Rating: {vul.get('risk_rating', '')}")
                new_doc.add_paragraph("Description")
                vul_index += 1
                continue

            corrected_text = correct_spelling(para.text)
            clean_paragraph_text = clean_text(corrected_text)
            new_paragraph = new_doc.add_paragraph(clean_paragraph_text)

            for run in para.runs:
                blips = run._element.findall('.//a:blip', namespaces={'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'})
                if blips:
                    blip = blips[0]
                    image_id = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                    if image_id and image_id in doc.part.related_parts:
                        image_part = doc.part.related_parts[image_id]

                        image_stream = io.BytesIO(image_part.blob)
                        image = Image.open(image_stream)

                        temp_image_path = os.path.join(temp_image_dir, f'image_{image_id}.png')
                        image.save(temp_image_path)

                        new_doc.add_picture(temp_image_path, width=Inches(3))

    new_doc.save(output_path)

    for image_file in os.listdir(temp_image_dir):
        os.remove(os.path.join(temp_image_dir, image_file))
    os.rmdir(temp_image_dir)

# Example usage
doc_path = r'C:\\Users\\FAYAZ PM\\PycharmProjects\\threatsview\\pythonProject\\scope of work test.doc'
docx_path = convert_doc_to_docx(doc_path)
remove_protection(docx_path)

start_heading = 'Detailed Findings'
end_heading = 'Annexure'
output_path = r'Esoutput.docx'
temp_image_dir = r'temp_images'

extract_content_between_headings(docx_path, start_heading, end_heading, output_path, temp_image_dir)
