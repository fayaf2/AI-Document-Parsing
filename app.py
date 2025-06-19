import os
import io
import re
import pythoncom
import tempfile
from PIL import Image
from flask import Flask, request, redirect, url_for, send_from_directory, render_template, jsonify
from werkzeug.utils import secure_filename
from docx import Document
from docx.shared import Inches
from spellchecker import SpellChecker
import win32com.client as win32

UPLOAD_FOLDER = 'static/uploads/'
ALLOWED_EXTENSIONS = {'docx', 'doc'}

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# Global variable to track progress
progress = 0

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def init_com():
    pythoncom.CoInitialize()

def convert_doc_to_docx(doc_path):
    init_com()
    word = win32.gencache.EnsureDispatch('Word.Application')
    doc = word.Documents.Open(doc_path)
    temp_docx_path = doc_path + 'x'  # Ensure .doc is converted to .docx
    doc.SaveAs(temp_docx_path, FileFormat=16)  # 16 is for docx
    doc.Close()
    word.Quit()
    return temp_docx_path

def remove_protection(docx_path):
    init_com()
    word = win32.gencache.EnsureDispatch('Word.Application')
    try:
        doc = word.Documents.Open(docx_path)
        if doc.ProtectionType != -1:
            doc.Unprotect()
        doc.Save()
        doc.Close()
    except Exception as e:
        print(f"Error: {e}")
    finally:
        word.Quit()

def extract_images_from_paragraph(para, temp_image_dir, new_doc):
    for run in para.runs:
        blips = run._element.findall('.//a:blip',
                                     namespaces={'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'})
        if blips:
            blip = blips[0]
            image_id = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
            if image_id and image_id in para.part.related_parts:
                image_part = para.part.related_parts[image_id]

                image_stream = io.BytesIO(image_part.blob)
                image = Image.open(image_stream)

                temp_image_path = os.path.join(temp_image_dir, f'image_{image_id}.png')
                image.save(temp_image_path)

                new_doc.add_picture(temp_image_path, width=Inches(3))

def clean_text(text):
    text = re.sub(r'\n+', '\n', text)
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'\s+([,.!?;:])', r'\1', text)
    text = re.sub(r'([,.!?;:])(\w)', r'\1 \2', text)
    return text

def correct_spelling(text):
    spell = SpellChecker()
    corrected_text = []
    for word in text.split():
        corrected_word = spell.correction(word)
        if corrected_word is None:
            corrected_word = word
        corrected_text.append(corrected_word)
    return ' '.join(corrected_text)

def extract_vulnerabilities(doc):
    vulnerabilities = []
    for table in doc.tables:
        vulnerability = {}
        for row in table.rows:
            if "Vulnerability Name" in row.cells[0].text:
                vulnerability_name = row.cells[1].text.strip()
                vulnerability["name"] = correct_spelling(vulnerability_name)
            elif "Affected Host" in row.cells[0].text:
                affected_host = row.cells[1].text.strip()
                vulnerability["host"] = correct_spelling(affected_host)
            elif "Risk Rating" in row.cells[0].text:
                risk_rating = row.cells[1].text.strip()
                vulnerability["risk_rating"] = correct_spelling(risk_rating)

        if vulnerability:
            vulnerabilities.append(vulnerability)

    return vulnerabilities

def extract_content_between_headings(doc_path, start_heading, output_path, temp_image_dir):
    global progress
    doc = Document(doc_path)
    new_doc = Document()

    vulnerabilities = extract_vulnerabilities(doc)

    in_section = False
    vul_index = 0
    start_heading_level = None
    subheadings = ["Description", "Severity", "Proof of Concept", "Remediations"]

    if not os.path.exists(temp_image_dir):
        os.makedirs(temp_image_dir)

    num_paragraphs = len(doc.paragraphs)

    # Find the level of the start_heading
    for para in doc.paragraphs:
        if para.style.name.startswith('Heading'):
            heading_text = para.text.strip()
            heading_level = int(re.search(r'\d+', para.style.name).group())

            if heading_text == start_heading:
                start_heading_level = heading_level
                in_section = True
                break

    if start_heading_level is None:
        raise ValueError("Could not find the start heading")

    # Process document to extract the content between headings
    for idx, para in enumerate(doc.paragraphs):
        if para.style.name.startswith('Heading'):
            heading_text = para.text.strip()
            heading_level = int(re.search(r'\d+', para.style.name).group())

            if heading_text == start_heading:
                in_section = True
            elif in_section and heading_level == start_heading_level:
                # Check if the next main heading contains any of the subheadings
                next_section_contains_subheadings = False
                for next_para in doc.paragraphs[idx + 1:]:
                    if next_para.style.name.startswith('Heading'):
                        next_heading_text = next_para.text.strip()
                        next_heading_level = int(re.search(r'\d+', next_para.style.name).group())
                        if next_heading_level == start_heading_level:
                            break
                    if any(subheading in next_para.text for subheading in subheadings):
                        next_section_contains_subheadings = True
                        break

                if not next_section_contains_subheadings:
                    in_section = False
                    continue

        if in_section:
            if para.text.strip() in subheadings and vul_index < len(vulnerabilities):
                vul = vulnerabilities[vul_index]
                new_doc.add_paragraph(f"Vulnerability Name: {vul.get('name', '')}")
                new_doc.add_paragraph(f"Affected Host: {vul.get('host', '')}")
                new_doc.add_paragraph(f"Risk Rating: {vul.get('risk_rating', '')}")
                new_doc.add_paragraph(para.text.strip())
                vul_index += 1
                continue

            corrected_text = correct_spelling(para.text)
            clean_paragraph_text = clean_text(corrected_text)
            new_paragraph = new_doc.add_paragraph(clean_paragraph_text)

            # Extract images within the paragraph only if within the section
            extract_images_from_paragraph(para, temp_image_dir, new_doc)

        # Update progress
        progress = int((idx / num_paragraphs) * 100)

    new_doc.save(output_path)

    for image_file in os.listdir(temp_image_dir):
        os.remove(os.path.join(temp_image_dir, image_file))
    os.rmdir(temp_image_dir)
    progress = 100  # Ensure progress reaches 100% at the end

@app.route('/')
def index():
    return render_template('upload.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    global progress
    if 'file' not in request.files:
        return redirect(request.url)
    file = request.files['file']
    if file.filename == '':
        return redirect(request.url)
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        temp_dir = tempfile.mkdtemp()
        file_path = os.path.join(temp_dir, filename)
        file.save(file_path)

        if filename.endswith('.doc'):
            file_path = convert_doc_to_docx(file_path)

        remove_protection(file_path)

        start_heading = 'Detailed Findings'
        output_filename = 'output.docx'
        output_path = os.path.join(app.config['UPLOAD_FOLDER'], output_filename)
        temp_image_dir = os.path.join(temp_dir, 'temp_images')

        # Reset progress before processing
        progress = 0
        extract_content_between_headings(file_path, start_heading, output_path, temp_image_dir)

        # Return URL to the processed file
        return redirect(url_for('view_file', filename=output_filename))

    return redirect(request.url)

@app.route('/view/<filename>')
def view_file(filename):
    return render_template('view_file.html', filename=filename)

@app.route('/uploads/<filename>')
def uploaded_file(filename):
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename)

@app.route('/progress', methods=['GET'])
def get_progress():
    global progress
    return jsonify({'progress': progress})

if __name__ == "__main__":
    app.run(debug=True)
